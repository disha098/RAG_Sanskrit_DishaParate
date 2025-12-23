from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.schema import Document as LCDocument
import os

DOC_PATH = "data/Rag-docs.docx"
VECTOR_DB_DIR = "vectorstore"

# Load DOCX
doc = Document(DOC_PATH)
full_text = []

for para in doc.paragraphs:
    text = para.text.strip()
    if len(text) > 20:   # ignore empty / very small lines
        full_text.append(text)

sanskrit_text = "\n".join(full_text)

# Chunk Sanskrit prose (larger chunks than verse-based texts)
splitter = RecursiveCharacterTextSplitter(
    chunk_size=600,
    chunk_overlap=120
)

chunks = splitter.split_text(sanskrit_text)

documents = [LCDocument(page_content=chunk) for chunk in chunks]

# Embeddings (works for Sanskrit + transliteration)
embeddings = HuggingFaceEmbeddings(
    model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2"
)

# Vector Store
db = FAISS.from_documents(documents, embeddings)
db.save_local(VECTOR_DB_DIR)

print("✅ Sanskrit DOCX dataset indexed successfully.")
